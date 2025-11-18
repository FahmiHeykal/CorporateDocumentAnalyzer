from docx import Document
import logging
from typing import List, Dict, Any

logger = logging.getLogger(__name__)

class DOCXExtractor:
    def __init__(self):
        pass
    
    def extract_text(self, file_path: str) -> str:
        """Ekstrak semua teks dari dokumen DOCX"""
        try:
            doc = Document(file_path)
            text_list = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(text_list)
            
            if not text:
                logger.warning(f"No extractable text found in {file_path}")
            
            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
            return ""  
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Ekstrak metadata dokumen DOCX"""
        metadata = {}
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata = {
                'title': core_props.title,
                'author': core_props.author,
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision,
                'paragraph_count': len(doc.paragraphs)
            }
        except Exception as e:
            logger.error(f"DOCX metadata extraction failed for {file_path}: {str(e)}")
        return metadata
    
    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """Ekstrak semua tabel dari dokumen DOCX"""
        tables = []
        try:
            doc = Document(file_path)
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
        except Exception as e:
            logger.error(f"DOCX table extraction failed for {file_path}: {str(e)}")
        return tables
    
    def extract_structure(self, file_path: str) -> Dict[str, List[str]]:
        """Ekstrak struktur dokumen: headings dan list"""
        structure = {'headings': [], 'lists': []}
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name
                if style_name.startswith('Heading') and paragraph.text.strip():
                    structure['headings'].append(paragraph.text)
                elif style_name == 'List Paragraph' and paragraph.text.strip():
                    structure['lists'].append(paragraph.text)
        except Exception as e:
            logger.error(f"DOCX structure extraction failed for {file_path}: {str(e)}")
        return structure
