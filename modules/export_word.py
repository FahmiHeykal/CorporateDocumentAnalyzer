from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
import io
import logging

logger = logging.getLogger(__name__)

class WordExporter:
    def __init__(self):
        pass
    
    def export(self, results: Dict, analysis_type: str) -> bytes:
        doc = Document()
        
        title = doc.add_heading('Corporate Document Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        if 'summary' in results:
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(results['summary'])
            doc.add_paragraph()
        
        if 'action_items' in results and results['action_items']:
            doc.add_heading('Action Items', level=1)
            for item in results['action_items'][:10]:
                p = doc.add_paragraph(item, style='List Bullet')
            doc.add_paragraph()
        
        if 'decisions' in results and results['decisions']:
            doc.add_heading('Decisions', level=1)
            for decision in results['decisions'][:10]:
                p = doc.add_paragraph(decision, style='List Bullet')
            doc.add_paragraph()
        
        if 'risks' in results and results['risks']:
            doc.add_heading('Risks', level=1)
            for risk in results['risks'][:10]:
                p = doc.add_paragraph(risk, style='List Bullet')
            doc.add_paragraph()
        
        if 'opportunities' in results and results['opportunities']:
            doc.add_heading('Opportunities', level=1)
            for opportunity in results['opportunities'][:10]:
                p = doc.add_paragraph(opportunity, style='List Bullet')
            doc.add_paragraph()
        
        if 'keywords' in results and results['keywords']:
            doc.add_heading('Key Keywords', level=1)
            keywords_text = ", ".join(results['keywords'][:15])
            doc.add_paragraph(keywords_text)
            doc.add_paragraph()
        
        if 'sentiment' in results:
            sentiment = results['sentiment']
            doc.add_heading('Sentiment Analysis', level=1)
            sentiment_text = f"Label: {sentiment.get('label', 'N/A')}, Score: {sentiment.get('score', 0):.2f}, Confidence: {sentiment.get('confidence', 0):.2f}"
            doc.add_paragraph(sentiment_text)
            doc.add_paragraph()
        
        if 'statistics' in results:
            stats = results['statistics']
            doc.add_heading('Document Statistics', level=1)
            
            stats_table = doc.add_table(rows=6, cols=2)
            stats_table.style = 'Light Grid Accent 1'
            
            headers = stats_table.rows[0].cells
            headers[0].text = 'Metric'
            headers[1].text = 'Value'
            
            data_rows = [
                ('Word Count', str(stats.get('word_count', 0))),
                ('Sentence Count', str(stats.get('sentence_count', 0))),
                ('Paragraph Count', str(stats.get('paragraph_count', 0))),
                ('Reading Time', f"{stats.get('reading_time_minutes', 0)} minutes"),
                ('Average Sentence Length', f"{stats.get('avg_sentence_length', 0):.2f} words")
            ]
            
            for i, (metric, value) in enumerate(data_rows, 1):
                cells = stats_table.rows[i].cells
                cells[0].text = metric
                cells[1].text = value
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()